"""
Servicio para generar documentos Word (DOCX) a partir de plantillas con sistema de placeholders.
Soporta placeholders en formato {{variable}} que se reemplazan automáticamente con los datos.
"""
import io
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
matplotlib.use('Agg')  # Backend sin interfaz gráfica
import matplotlib.pyplot as plt

TEMPLATE_PATH = Path(__file__).parent.parent / "mc-tipo.docx"


def _get_nested_value(data: Dict[str, Any], path: str) -> Optional[Any]:
    """
    Obtiene un valor anidado de un diccionario usando notación de puntos.

    Ejemplo: _get_nested_value(data, "seismic.result.Qbasx")
    retorna data["seismic"]["result"]["Qbasx"]
    """
    keys = path.split(".")
    value = data

    for key in keys:
        if isinstance(value, dict):
            value = value.get(key)
            if value is None:
                return None
        else:
            return None

    return value


def _format_value(value: Any, decimal_places: int = 2) -> str:
    """Formatea un valor para mostrar en el documento."""
    if value is None:
        return ""

    if isinstance(value, (int, float)):
        return f"{value:.{decimal_places}f}"

    return str(value)


def _build_context(data: Dict[str, Any], project_name: str) -> Dict[str, Any]:
    """
    Construye el contexto con todas las variables disponibles para reemplazo.
    Este es el diccionario completo de placeholders disponibles.
    """
    context = {
        # Información del proyecto
        "projectName": project_name,
        "currentDate": datetime.now().strftime("%d de %B de %Y"),

        # Descripción del edificio (si existe)
        "buildingDescription": data.get("buildingDescription", {}).get("text", ""),
        "buildingLocation": data.get("buildingDescription", {}).get("location", ""),
        "buildingArea": data.get("buildingDescription", {}).get("area", ""),
        "buildingHeight": data.get("buildingDescription", {}).get("height", ""),
    }

    # CARGAS VIVAS
    if "liveLoad" in data and data["liveLoad"]:
        ll = data["liveLoad"]
        context.update({
            "liveLoad.buildingType": ll.get("buildingType", ""),
            "liveLoad.usage": ll.get("usage", ""),
            "liveLoad.uniformLoad": _format_value(ll.get("uniformLoad")),
            "liveLoad.uniformLoadRaw": ll.get("uniformLoadRaw", ""),
            "liveLoad.concentratedLoad": _format_value(ll.get("concentratedLoad")),
            "liveLoad.concentratedLoadRaw": ll.get("concentratedLoadRaw", ""),
        })

    # REDUCCIÓN DE CARGAS VIVAS
    if "reduction" in data and data["reduction"]:
        red = data["reduction"]
        context.update({
            "reduction.elementType": red.get("elementType", ""),
            "reduction.tributaryArea": _format_value(red.get("tributaryArea")),
            "reduction.baseLoad": _format_value(red.get("baseLoad")),
            "reduction.reducedLoad": _format_value(red.get("reducedLoad")),
        })

    # VIENTO
    if "wind" in data and data["wind"]:
        wind = data["wind"]
        context.update({
            "wind.environment": wind.get("environment", ""),
            "wind.height": _format_value(wind.get("height")),
            "wind.q": _format_value(wind.get("q")),
            "wind.message": wind.get("message", ""),
        })

    # NIEVE
    if "snow" in data and data["snow"]:
        snow = data["snow"]
        context.update({
            "snow.latitudeBand": snow.get("latitudeBand", ""),
            "snow.altitudeBand": snow.get("altitudeBand", ""),
            "snow.thermalCondition": snow.get("thermalCondition", ""),
            "snow.importanceCategory": snow.get("importanceCategory", ""),
            "snow.exposureCategory": snow.get("exposureCategory", ""),
            "snow.exposureCondition": snow.get("exposureCondition", ""),
            "snow.surfaceType": snow.get("surfaceType", ""),
            "snow.roofPitch": _format_value(snow.get("roofPitch")),
            "snow.pg": _format_value(snow.get("pg"), 3),
            "snow.ct": _format_value(snow.get("ct"), 3),
            "snow.ce": _format_value(snow.get("ce"), 3),
            "snow.I": _format_value(snow.get("I"), 3),
            "snow.cs": _format_value(snow.get("cs"), 3),
            "snow.pf": _format_value(snow.get("pf"), 3),
        })

    # SISMO
    if "seismic" in data and data["seismic"]:
        params = data["seismic"].get("params", {})
        result = data["seismic"].get("result", {})

        context.update({
            # Parámetros de entrada
            "seismic.params.category": params.get("category", ""),
            "seismic.params.zone": params.get("zone", ""),
            "seismic.params.soil": params.get("soil", ""),
            "seismic.params.rs": _format_value(params.get("rs")),
            "seismic.params.ps": _format_value(params.get("ps")),
            "seismic.params.tx": _format_value(params.get("tx"), 3),
            "seismic.params.ty": _format_value(params.get("ty"), 3),
            "seismic.params.r0": _format_value(params.get("r0")),

            # Resultados
            "seismic.result.intensityFactor": _format_value(result.get("intensityFactor"), 3),
            "seismic.result.zoneFactor": _format_value(result.get("zoneFactor"), 3),
            "seismic.result.CMax": _format_value(result.get("CMax", result.get("C_max")), 3),
            "seismic.result.CMin": _format_value(result.get("CMin", result.get("C_min")), 3),
            "seismic.result.Q0x": _format_value(result.get("Q0x")),
            "seismic.result.Q0y": _format_value(result.get("Q0y")),
            "seismic.result.Q0Min": _format_value(result.get("Q0Min", result.get("Q0_min"))),
            "seismic.result.Q0Max": _format_value(result.get("Q0Max", result.get("Q0_max"))),
            "seismic.result.Qbasx": _format_value(result.get("Qbasx")),
            "seismic.result.Qbasy": _format_value(result.get("Qbasy")),
        })

    # CÁLCULOS ESTRUCTURALES
    if "structural" in data and data["structural"]:
        struct = data["structural"]

        # Pilar de Hormigón Armado
        if "concreteColumn" in struct and struct["concreteColumn"]:
            cc = struct["concreteColumn"]
            context.update({
                "concrete.column.axialCapacity": _format_value(cc.get("axialCapacity")),
                "concrete.column.axialCapacityRatio": _format_value(cc.get("axialCapacityRatio"), 3),
                "concrete.column.longitudinalSteel.numBars": str(cc.get("longitudinalSteel", {}).get("numBars", "")),
                "concrete.column.longitudinalSteel.barDiameter": str(cc.get("longitudinalSteel", {}).get("barDiameter", "")),
                "concrete.column.longitudinalSteel.totalArea": _format_value(cc.get("longitudinalSteel", {}).get("totalArea")),
                "concrete.column.longitudinalSteel.ratio": _format_value(cc.get("longitudinalSteel", {}).get("ratio"), 4),
                "concrete.column.transverseSteel.diameter": str(cc.get("transverseSteel", {}).get("diameter", "")),
                "concrete.column.transverseSteel.spacing": _format_value(cc.get("transverseSteel", {}).get("spacing"), 0),
                "concrete.column.shearCapacityRatioX": _format_value(cc.get("shearCapacityRatioX"), 3),
                "concrete.column.shearCapacityRatioY": _format_value(cc.get("shearCapacityRatioY"), 3),
                "concrete.column.slendernessRatio": _format_value(cc.get("slendernessRatio")),
                "concrete.column.magnificationFactor": _format_value(cc.get("magnificationFactor"), 3),
                "concrete.column.isSlender": str(cc.get("isSlender", "")),
            })

        # Viga de Hormigón Armado
        if "concreteBeam" in struct and struct["concreteBeam"]:
            cb = struct["concreteBeam"]
            context.update({
                "concrete.beam.positiveReinforcement.numBars": str(cb.get("positiveReinforcement", {}).get("numBars", "")),
                "concrete.beam.positiveReinforcement.barDiameter": str(cb.get("positiveReinforcement", {}).get("barDiameter", "")),
                "concrete.beam.positiveReinforcement.totalArea": _format_value(cb.get("positiveReinforcement", {}).get("totalArea")),
                "concrete.beam.positiveReinforcement.ratio": _format_value(cb.get("positiveReinforcement", {}).get("ratio"), 4),
                "concrete.beam.negativeReinforcement.numBars": str(cb.get("negativeReinforcement", {}).get("numBars", "")),
                "concrete.beam.negativeReinforcement.barDiameter": str(cb.get("negativeReinforcement", {}).get("barDiameter", "")),
                "concrete.beam.negativeReinforcement.totalArea": _format_value(cb.get("negativeReinforcement", {}).get("totalArea")),
                "concrete.beam.negativeReinforcement.ratio": _format_value(cb.get("negativeReinforcement", {}).get("ratio"), 4),
                "concrete.beam.transverseSteel.diameter": str(cb.get("transverseSteel", {}).get("diameter", "")),
                "concrete.beam.transverseSteel.spacing": _format_value(cb.get("transverseSteel", {}).get("spacing"), 0),
                "concrete.beam.shearCapacityRatio": _format_value(cb.get("shearCapacityRatio"), 3),
                "concrete.beam.deflectionCheck": cb.get("deflectionCheck", ""),
                "concrete.beam.effectiveDepth": _format_value(cb.get("effectiveDepth")),
            })

        # Pilar de Acero
        if "steelColumn" in struct and struct["steelColumn"]:
            sc = struct["steelColumn"]
            axial_capacity = _format_value(sc.get("axialCapacity"))
            axial_ratio = _format_value(sc.get("axialCapacityRatio"), 3)
            moment_capacity_x = _format_value(sc.get("momentCapacityX"))
            moment_capacity_y = _format_value(sc.get("momentCapacityY"))
            moment_ratio_x = _format_value(sc.get("momentCapacityRatioX"), 3)
            moment_ratio_y = _format_value(sc.get("momentCapacityRatioY"), 3)
            slender_parameter = _format_value(sc.get("slendernessParameter", sc.get("slendernessMax")))
            passes_str = str(sc.get("checkPasses", sc.get("passes", "")))

            context.update({
                "steel.column.section": sc.get("section", ""),
                "steel.column.axialCapacity": axial_capacity,
                "steel.column.pn": axial_capacity,
                "steel.column.axialCapacityRatio": axial_ratio,
                "steel.column.axialRatio": axial_ratio,
                "steel.column.momentCapacityX": moment_capacity_x,
                "steel.column.mnX": moment_capacity_x,
                "steel.column.momentCapacityY": moment_capacity_y,
                "steel.column.mnY": moment_capacity_y,
                "steel.column.momentCapacityRatioX": moment_ratio_x,
                "steel.column.flexureRatioX": moment_ratio_x,
                "steel.column.momentCapacityRatioY": moment_ratio_y,
                "steel.column.flexureRatioY": moment_ratio_y,
                "steel.column.slendernessX": _format_value(sc.get("slendernessX")),
                "steel.column.slendernessY": _format_value(sc.get("slendernessY")),
                "steel.column.slendernessMax": _format_value(sc.get("slendernessMax")),
                "steel.column.slendernessParameter": slender_parameter,
                "steel.column.lambdaC": slender_parameter,
                "steel.column.interactionRatio": _format_value(sc.get("interactionRatio"), 3),
                "steel.column.checkPasses": passes_str,
                "steel.column.passes": passes_str,
                "steel.column.checkStatus": sc.get("checkStatus", ""),
            })

        # Viga de Acero
        if "steelBeam" in struct and struct["steelBeam"]:
            sb = struct["steelBeam"]
            moment_capacity = _format_value(sb.get("momentCapacity"))
            moment_ratio = _format_value(sb.get("momentCapacityRatio"), 3)
            shear_capacity = _format_value(sb.get("shearCapacity"))
            shear_ratio = _format_value(sb.get("shearCapacityRatio"), 3)
            passes_str = str(sb.get("checkPasses", sb.get("passes", "")))

            context.update({
                "steel.beam.section": sb.get("section", ""),
                "steel.beam.momentCapacity": moment_capacity,
                "steel.beam.mn": moment_capacity,
                "steel.beam.momentCapacityRatio": moment_ratio,
                "steel.beam.flexureRatio": moment_ratio,
                "steel.beam.shearCapacity": shear_capacity,
                "steel.beam.vn": shear_capacity,
                "steel.beam.shearCapacityRatio": shear_ratio,
                "steel.beam.shearRatio": shear_ratio,
                "steel.beam.deflection": _format_value(sb.get("deflection")),
                "steel.beam.deflectionLimit": _format_value(sb.get("deflectionLimit")),
                "steel.beam.deflectionRatio": _format_value(sb.get("deflectionRatio"), 3),
                "steel.beam.lateralBracingLength": _format_value(sb.get("lateralBracingLength"), 0),
                "steel.beam.checkPasses": passes_str,
                "steel.beam.passes": passes_str,
                "steel.beam.checkStatus": sb.get("checkStatus", ""),
            })

        # Pilar de Madera
        if "woodColumn" in struct and struct["woodColumn"]:
            wc = struct["woodColumn"]
            context.update({
                "wood.column.woodType": wc.get("woodType", ""),
                "wood.column.area": _format_value(wc.get("area")),
                "wood.column.axialCapacity": _format_value(wc.get("axialCapacity")),
                "wood.column.axialCapacityRatio": _format_value(wc.get("axialCapacityRatio"), 3),
                "wood.column.slendernessX": _format_value(wc.get("slendernessX")),
                "wood.column.slendernessY": _format_value(wc.get("slendernessY")),
                "wood.column.slendernessMax": _format_value(wc.get("slendernessMax")),
                "wood.column.stabilityFactor": _format_value(wc.get("stabilityFactor"), 3),
                "wood.column.isSlender": str(wc.get("isSlender", "")),
                "wood.column.allowableStress": _format_value(wc.get("allowableStress")),
                "wood.column.checkStatus": wc.get("checkStatus", ""),
            })


            # Aliases para compatibilidad con plantilla
            axial_capacity = _format_value(wc.get("axialCapacity"))
            axial_ratio = _format_value(wc.get("axialCapacityRatio"), 3)
            passes_str = str(wc.get("checkPasses", wc.get("passes", "")))
            context.update({
                "wood.column.pn": axial_capacity,
                "wood.column.utilizationRatio": axial_ratio,
                "wood.column.checkPasses": passes_str,
                "wood.column.passes": passes_str,
            })
        # Viga de Madera
        if "woodBeam" in struct and struct["woodBeam"]:
            wb = struct["woodBeam"]
            context.update({
                "wood.beam.woodType": wb.get("woodType", ""),
                "wood.beam.section": wb.get("section", ""),
                "wood.beam.area": _format_value(wb.get("area")),
                "wood.beam.sectionModulus": _format_value(wb.get("sectionModulus")),
                "wood.beam.momentOfInertia": _format_value(wb.get("momentOfInertia")),
                "wood.beam.nominalMomentCapacity": _format_value(wb.get("nominalMomentCapacity", wb.get("momentCapacity"))),
                "wood.beam.nominalShearCapacity": _format_value(wb.get("nominalShearCapacity", wb.get("shearCapacity"))),
                "wood.beam.utilization": _format_value(wb.get("utilization", wb.get("overallUtilization")), 3),
                "wood.beam.flexureStress": _format_value(wb.get("flexureStress")),
                "wood.beam.allowableFlexureStress": _format_value(wb.get("allowableFlexureStress")),
                "wood.beam.flexureRatio": _format_value(wb.get("flexureRatio"), 3),
                "wood.beam.shearStress": _format_value(wb.get("shearStress")),
                "wood.beam.allowableShearStress": _format_value(wb.get("allowableShearStress")),
                "wood.beam.shearRatio": _format_value(wb.get("shearRatio"), 3),
                "wood.beam.deflection": _format_value(wb.get("deflection")),
                "wood.beam.deflectionLimit": _format_value(wb.get("deflectionLimit")),
                "wood.beam.deflectionRatio": _format_value(wb.get("deflectionRatio"), 3),
                "wood.beam.lateralStabilityFactor": _format_value(wb.get("lateralStabilityFactor"), 3),
                "wood.beam.checkPasses": str(wb.get("checkPasses", wb.get("passes", ""))),
                "wood.beam.checkStatus": wb.get("checkStatus", ""),
            })


            # Aliases para compatibilidad con plantilla
            nominal_moment = _format_value(wb.get("nominalMomentCapacity", wb.get("momentCapacity")))
            nominal_shear = _format_value(wb.get("nominalShearCapacity", wb.get("shearCapacity")))
            utilization = _format_value(wb.get("utilization", wb.get("overallUtilization")), 3)
            passes_str = str(wb.get("checkPasses", wb.get("passes", "")))
            context.update({
                "wood.beam.mn": nominal_moment,
                "wood.beam.vn": nominal_shear,
                "wood.beam.utilizationRatio": utilization,
                "wood.beam.passes": passes_str,
                "wood.beam.checkPasses": passes_str,
            })
        # Zapata
        if "footing" in struct and struct["footing"]:
            ft = struct["footing"]
            context.update({
                "footing.footingType": ft.get("footingType", ""),
                "footing.dimensions.length": _format_value(ft.get("dimensions", {}).get("length"), 2),
                "footing.dimensions.width": _format_value(ft.get("dimensions", {}).get("width"), 2),
                "footing.dimensions.depth": _format_value(ft.get("dimensions", {}).get("depth"), 1),
                "footing.dimensions.area": _format_value(ft.get("dimensions", {}).get("area"), 2),
                "footing.soilPressures.max": _format_value(ft.get("soilPressures", {}).get("max"), 2),
                "footing.soilPressures.min": _format_value(ft.get("soilPressures", {}).get("min"), 2),
                "footing.soilPressures.average": _format_value(ft.get("soilPressures", {}).get("average"), 2),
                "footing.soilPressures.ratio": _format_value(ft.get("soilPressures", {}).get("ratio"), 3),
                "footing.lateralPressures.static": _format_value(ft.get("lateralPressures", {}).get("static")),
                "footing.lateralPressures.dynamic": _format_value(ft.get("lateralPressures", {}).get("dynamic")),
                "footing.lateralPressures.seismic": _format_value(ft.get("lateralPressures", {}).get("seismic")),
                "footing.lateralPressures.total": _format_value(ft.get("lateralPressures", {}).get("total")),
                "footing.punchingShear.appliedForce": _format_value(ft.get("punchingShear", {}).get("appliedForce")),
                "footing.punchingShear.capacity": _format_value(ft.get("punchingShear", {}).get("capacity")),
                "footing.punchingShear.ratio": _format_value(ft.get("punchingShear", {}).get("ratio"), 3),
                "footing.punchingShear.criticalPerimeter": _format_value(ft.get("punchingShear", {}).get("criticalPerimeter"), 1),
                "footing.flexuralShear.appliedForce": _format_value(ft.get("flexuralShear", {}).get("appliedForce")),
                "footing.flexuralShear.capacity": _format_value(ft.get("flexuralShear", {}).get("capacity")),
                "footing.flexuralShear.ratio": _format_value(ft.get("flexuralShear", {}).get("ratio"), 3),
                "footing.checkPasses": str(ft.get("checkPasses", ft.get("passes", ""))),
                "footing.checkStatus": ft.get("checkStatus", ""),
            })

            # Refuerzo (varía según el tipo)
            reinforcement = ft.get("reinforcement", {})
            if "xDirection" in reinforcement:  # Zapata aislada
                xdir = reinforcement.get("xDirection", {})
                ydir = reinforcement.get("yDirection", {})
                context.update({
                    "footing.reinforcement.longitudinalSteel": _format_value(xdir.get("totalArea", xdir.get("area"))),
                    "footing.reinforcement.transverseSteel": _format_value(ydir.get("totalArea", ydir.get("area"))),
                    "footing.reinforcement.xDirection.barDiameter": str(xdir.get("barDiameter", "")),
                    "footing.reinforcement.xDirection.spacing": _format_value(xdir.get("spacing"), 1),
                    "footing.reinforcement.yDirection.barDiameter": str(ydir.get("barDiameter", "")),
                    "footing.reinforcement.yDirection.spacing": _format_value(ydir.get("spacing"), 1),
                })
            elif "main" in reinforcement:  # Zapata corrida
                main = reinforcement.get("main", {})
                dist = reinforcement.get("distribution", {})
                context.update({
                    "footing.reinforcement.longitudinalSteel": _format_value(main.get("totalArea", main.get("area"))),
                    "footing.reinforcement.transverseSteel": _format_value(dist.get("totalArea", dist.get("area"))),
                    "footing.reinforcement.main.barDiameter": str(main.get("barDiameter", "")),
                    "footing.reinforcement.main.spacing": _format_value(main.get("spacing"), 1),
                    "footing.reinforcement.main.direction": main.get("direction", ""),
                    "footing.reinforcement.distribution.barDiameter": str(dist.get("barDiameter", "")),
                    "footing.reinforcement.distribution.spacing": _format_value(dist.get("spacing"), 1),
                    "footing.reinforcement.distribution.direction": dist.get("direction", ""),
                })

            # Aliases para compatibilidad con la plantilla existente
            dims = ft.get("dimensions", {})
            soil = ft.get("soilPressures", {})
            punching = ft.get("punchingShear", {})
            flex = ft.get("flexuralShear", {})
            length = _format_value(dims.get("length"), 2)
            width = _format_value(dims.get("width"), 2)
            depth = _format_value(dims.get("depth"), 1)
            as_long = context.get("footing.reinforcement.longitudinalSteel", "")
            as_trans = context.get("footing.reinforcement.transverseSteel", "")
            bar_diam = (
                str(ft.get("reinforcement", {}).get("xDirection", {}).get("barDiameter"))
                or str(ft.get("reinforcement", {}).get("main", {}).get("barDiameter", ""))
            )
            spacing = (
                _format_value(ft.get("reinforcement", {}).get("xDirection", {}).get("spacing"), 1)
                or _format_value(ft.get("reinforcement", {}).get("main", {}).get("spacing"), 1)
            )
            context.update({
                "footing.length": length,
                "footing.width": width,
                "footing.depth": depth,
                "footing.soilPressureMax": _format_value(soil.get("max"), 2),
                "footing.soilPressureMin": _format_value(soil.get("min"), 2),
                "footing.punchingShearRatio": _format_value(punching.get("ratio"), 3),
                "footing.beamShearRatio": _format_value(flex.get("ratio"), 3),
                "footing.asLongitudinal": as_long,
                "footing.asTransverse": as_trans,
                "footing.barDiameter": bar_diam,
                "footing.spacing": spacing,
                "footing.passes": str(ft.get("checkPasses", ft.get("passes", ""))),
            })

    # Permitir inyección de placeholders de tablas y extras
    tables = data.get("tables")
    if isinstance(tables, dict):
        context.update(tables)

    extra_placeholders = (
        data.get("placeholders")
        or data.get("extraPlaceholders")
        or data.get("additionalPlaceholders")
    )
    if isinstance(extra_placeholders, dict):
        context.update(extra_placeholders)

    return context


def _replace_placeholders_in_text(text: str, context: Dict[str, Any]) -> str:
    """
    Reemplaza todos los placeholders {{variable}} en un texto con sus valores del contexto.
    """
    def replacer(match):
        placeholder = match.group(1).strip()
        value = context.get(placeholder, f"{{{{placeholder}}}}")  # Mantiene el placeholder si no existe
        return str(value) if value else ""

    # Patrón para encontrar {{variable}}
    pattern = r'\{\{([^}]+)\}\}'
    return re.sub(pattern, replacer, text)


def _replace_placeholders_in_runs(paragraph, context: Dict[str, Any]):
    """
    Reemplaza placeholders en los runs de un párrafo.

    Esta función maneja el caso donde un placeholder puede estar dividido
    entre múltiples runs de texto en Word, lo que sucede frecuentemente cuando
    se edita el documento manualmente.
    """
    # Obtener el texto completo del párrafo
    full_text = paragraph.text

    # Si no hay placeholders, no hacemos nada
    if '{{' not in full_text or '}}' not in full_text:
        return

    # Patrón para encontrar placeholders
    pattern = r'\{\{([^}]+)\}\}'

    # Encontrar todos los placeholders y sus posiciones en el texto completo
    matches = list(re.finditer(pattern, full_text))

    if not matches:
        return

    # Construir el nuevo texto reemplazando los placeholders de atrás hacia adelante
    new_text = full_text
    for match in reversed(matches):
        placeholder = match.group(1).strip()
        # Buscar el valor en el contexto
        value = context.get(placeholder, "")

        # Si no se encuentra el valor, mantener el placeholder para debug
        if not value and placeholder not in context:
            value = f"{{{{placeholder}}}}"

        # Reemplazar el placeholder con su valor
        new_text = new_text[:match.start()] + str(value) + new_text[match.end():]

    # Método más robusto: mantener el formato del primer run
    # y poner todo el texto allí
    if paragraph.runs:
        # Guardar el formato del primer run
        first_run = paragraph.runs[0]

        # Limpiar todos los runs
        for run in paragraph.runs:
            run.text = ""

        # Poner el nuevo texto en el primer run
        first_run.text = new_text
    else:
        # Si no hay runs, crear uno nuevo
        paragraph.add_run(new_text)


def _replace_placeholders_in_paragraphs(doc: Document, context: Dict[str, Any]):
    """Reemplaza placeholders en todos los párrafos del documento preservando el formato."""
    for para in doc.paragraphs:
        if '{{' in para.text and '}}' in para.text:
            _replace_placeholders_in_runs(para, context)


def _replace_placeholders_in_tables(doc: Document, context: Dict[str, Any]):
    """Reemplaza placeholders en todas las tablas del documento preservando el formato."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '{{' in para.text and '}}' in para.text:
                        _replace_placeholders_in_runs(para, context)


def _generate_seismic_spectrum_chart(data: Dict[str, Any]) -> Optional[io.BytesIO]:
    """
    Genera un gráfico de espectros sísmicos (Aceleración vs Período).

    Returns:
        BytesIO con la imagen del gráfico en formato PNG, o None si no hay datos
    """
    if "seismic" not in data or "seismic" not in data:
        return None

    seismic_data = data["seismic"]
    if "result" not in seismic_data or "spectrum" not in seismic_data["result"]:
        return None

    spectrum = seismic_data["result"]["spectrum"]
    if not spectrum:
        return None

    # Extraer datos para el gráfico
    periods = [point["period"] for point in spectrum]
    sa_x = [point.get("SaX", point.get("Sa_x", 0)) for point in spectrum]
    sa_y = [point.get("SaY", point.get("Sa_y", 0)) for point in spectrum]

    # Crear el gráfico
    plt.figure(figsize=(10, 6))
    plt.plot(periods, sa_x, 'b-', linewidth=2, label='Espectro X')
    plt.plot(periods, sa_y, 'r--', linewidth=2, label='Espectro Y')
    plt.xlabel('Período T (s)', fontsize=12)
    plt.ylabel('Aceleración espectral Sa (g)', fontsize=12)
    plt.title('Espectros de Diseño Sísmico', fontsize=14, fontweight='bold')
    plt.grid(True, alpha=0.3)
    plt.legend(fontsize=10)
    plt.tight_layout()

    # Guardar en buffer
    buffer = io.BytesIO()
    plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buffer.seek(0)

    return buffer


def _insert_spectrum_chart(doc: Document, data: Dict[str, Any]):
    """
    Inserta el gráfico de espectros sísmicos en el documento donde encuentre {{spectrumChart}}.
    """
    chart_buffer = _generate_seismic_spectrum_chart(data)
    if not chart_buffer:
        return

    # Buscar el placeholder {{spectrumChart}} en párrafos
    for para in doc.paragraphs:
        if '{{spectrumChart}}' in para.text:
            # Limpiar el texto del párrafo
            para.text = para.text.replace('{{spectrumChart}}', '')
            # Insertar la imagen
            run = para.add_run()
            run.add_picture(chart_buffer, width=Inches(6))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break


def generate_design_base_document(data: Dict[str, Any], project_name: str = "Proyecto") -> bytes:
    """
    Genera un documento Word a partir de la plantilla usando sistema de placeholders.

    Los placeholders se escriben en el Word como {{variable}} y se reemplazan automáticamente.

    Args:
        data: Diccionario con los datos de bases de cálculo
        project_name: Nombre del proyecto

    Returns:
        Bytes del documento Word generado
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Plantilla no encontrada: {TEMPLATE_PATH}")

    # Cargar plantilla
    doc = Document(str(TEMPLATE_PATH))

    # Construir contexto con todas las variables
    context = _build_context(data, project_name)

    # Reemplazar placeholders en todo el documento
    _replace_placeholders_in_paragraphs(doc, context)
    _replace_placeholders_in_tables(doc, context)

    # Insertar gráfico de espectros si existe el placeholder
    _insert_spectrum_chart(doc, data)

    # Guardar en buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()
