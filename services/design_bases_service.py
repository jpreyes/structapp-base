import csv
import io
import json
import math
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from api.schemas.design_bases import DesignBaseExportPayload

DATA_DIR = Path(__file__).resolve().parent.parent / "api" / "data"
with (DATA_DIR / "live_loads.json").open(encoding="utf-8") as fp:
    _LIVE_LOADS_RAW: Dict[str, Dict[str, Dict[str, str]]] = json.load(fp)


def _safe_float(value: str) -> Optional[float]:
    try:
        return float(value.replace(",", "."))
    except ValueError:
        return None


LIVE_LOADS: Dict[str, List[Dict[str, object]]] = {
    category: [
        {
            "usage": usage,
            "uniform_load": _safe_float(data["uniform_load"]),
            "uniform_load_raw": data["uniform_load"],
            "concentrated_load": _safe_float(data["concentrated_load"]),
            "concentrated_load_raw": data["concentrated_load"],
        }
        for usage, data in usages.items()
    ]
    for category, usages in _LIVE_LOADS_RAW.items()
}

LIVE_LOAD_LOOKUP: Dict[Tuple[str, str], Dict[str, object]] = {
    (category, entry["usage"]): entry
    for category, entries in LIVE_LOADS.items()
    for entry in entries
}

KLL_FACTORS: Dict[str, float] = {
    "Columnas interiores": 4.0,
    "Columnas exteriores sin losas en voladizo": 4.0,
    "Columnas de borde con losas en voladizos": 3.0,
    "Columnas de esquina con losas en voladizo": 2.0,
    "Vigas de borde sin losas en voladizo": 2.0,
    "Vigas interiores": 2.0,
    "Vigas de borde con losas en voladizo": 1.0,
    "Vigas en voladizo": 1.0,
    "Losas en una dirección": 1.0,
    "Losas en dos direcciones": 1.0,
    "Elementos sin transferencia continua de corte": 1.0,
}

WIND_PROFILES: Dict[str, List[Tuple[float, float, float, float]]] = {
    "Construcciones en ciudad o similar": [
        (0.0, 15.0, 0.2 / 15.0, 0.55),
        (15.0, 20.0, 0.1 / 5.0, 0.75),
        (20.0, 30.0, 0.1 / 10.0, 0.85),
        (30.0, 40.0, 0.08 / 10.0, 0.95),
        (40.0, 50.0, 0.05 / 10.0, 1.03),
        (50.0, 75.0, 0.13 / 25.0, 1.08),
        (75.0, 100.0, 0.10 / 25.0, 1.21),
        (100.0, 150.0, 0.18 / 50.0, 1.31),
        (150.0, 200.0, 0.13 / 50.0, 1.49),
        (200.0, 300.0, 0.24 / 100.0, 1.62),
    ],
    "Terrenos abiertos o sin obstáculos": [
        (0.0, 4.0, 0.0, 0.7),
        (4.0, 7.0, 0.25 / 3.0, 0.7),
        (7.0, 10.0, 0.11 / 3.0, 0.95),
        (10.0, 15.0, 0.12 / 5.0, 1.06),
        (15.0, 20.0, 0.08 / 5.0, 1.18),
        (20.0, 30.0, 0.12 / 10.0, 1.26),
        (30.0, 40.0, 0.10 / 10.0, 1.38),
        (40.0, 50.0, 0.09 / 10.0, 1.48),
        (50.0, 60.0, 0.08 / 10.0, 1.57),
        (60.0, 70.0, 0.07 / 10.0, 1.65),
        (70.0, 80.0, 0.08 / 10.0, 1.72),
        (80.0, 90.0, 0.08 / 10.0, 1.8),
        (90.0, 120.0, 0.15 / 30.0, 1.88),
        (120.0, 150.0, 0.10 / 30.0, 2.03),
    ],
}

SNOW_LOAD_MAP = json.loads((DATA_DIR / "snow_map.json").read_text(encoding="utf-8"))

THERMAL_FACTORS = {
    "Todas las estructuras, excepto las indicadas": 1.0,
    "Estructuras justo sobre el punto de congelamiento": 1.1,
    "Estructuras no calefaccionadas": 1.2,
    "Estructuras calefaccionadas en forma excepcional": 0.85,
}

IMPORTANCE_FACTORS = {
    "Categoría I": 0.8,
    "Categoría II": 1.0,
    "Categoría III": 1.1,
    "Categoría IV": 1.2,
}

EXPOSURE_FACTORS = {
    "Categoría B: áreas urbanas, suburbanas y boscosas...": {
        "Totalmente expuesto": 0.9,
        "Parcialmente expuesto": 1.0,
        "A sotavento": 1.2,
    },
    "Categoría C: terrenos abiertos, con obstáculos dispersos...": {
        "Totalmente expuesto": 0.9,
        "Parcialmente expuesto": 1.0,
        "A sotavento": 1.1,
    },
    "Categoría D: planicies, áreas sin obstáculos...": {
        "Totalmente expuesto": 0.8,
        "Parcialmente expuesto": 0.9,
        "A sotavento": 1.0,
    },
    "Otra": {
        "Totalmente expuesto": 0.7,
        "Parcialmente expuesto": 0.8,
        "A sotavento": 1.0,
    },
}

SURFACE_OPTIONS = {"Superficie lisas sin obstáculos", "Todas las otras superficies"}

SEISMIC_SOIL = {
    "Suelo A": {"S_S": 0.9, "T0_S": 0.15, "TP_S": 0.2, "n_S": 1.0, "p_S": 2.0},
    "Suelo B": {"S_S": 1.0, "T0_S": 0.3, "TP_S": 0.35, "n_S": 1.33, "p_S": 1.5},
    "Suelo C": {"S_S": 1.05, "T0_S": 0.4, "TP_S": 0.45, "n_S": 1.4, "p_S": 1.6},
    "Suelo D": {"S_S": 1.2, "T0_S": 0.75, "TP_S": 0.85, "n_S": 1.8, "p_S": 1.0},
    "Suelo E": {"S_S": 1.3, "T0_S": 1.2, "TP_S": 1.35, "n_S": 1.8, "p_S": 1.0},
}


def list_live_load_categories() -> Dict[str, List[Dict[str, object]]]:
    return LIVE_LOADS


def get_live_load(category: str, usage: str) -> Dict[str, object]:
    key = (category, usage)
    if key not in LIVE_LOAD_LOOKUP:
        raise KeyError(f"No existen datos para {category} / {usage}")
    return LIVE_LOAD_LOOKUP[key]


def calculate_live_load_reduction(element_type: str, tributary_area: float, base_load: float) -> float:
    if element_type not in KLL_FACTORS:
        raise KeyError(f"Tipo de elemento no reconocido: {element_type}")
    if tributary_area <= 0 or base_load <= 0:
        raise ValueError("El área tributaria y la carga base deben ser mayores que cero.")
    factor = KLL_FACTORS[element_type]
    reduction = base_load * (0.25 + 4.75 / math.sqrt(factor * tributary_area))
    return min(reduction, base_load)


def interpolate_profile(height: float, profile: List[Tuple[float, float, float, float]]) -> Optional[float]:
    for lower, upper, slope, base in profile:
        if lower <= height < upper:
            return slope * (height - lower) + base
    return None


def calculate_wind_pressure(environment: str, height: float) -> Dict[str, Optional[float]]:
    if environment not in WIND_PROFILES:
        raise KeyError(f"Ambiente no soportado: {environment}")
    if height <= 0:
        raise ValueError("La altura debe ser positiva.")
    profile = WIND_PROFILES[environment]
    q = interpolate_profile(height, profile)
    if q is None:
        return {"q": None, "message": "Fuera del rango tabulado. Revisa NCh432."}
    return {"q": round(q, 4), "message": None}


def lookup_snow_load(latitude_band: str, altitude_band: str) -> str:
    try:
        alt_map = SNOW_LOAD_MAP[latitude_band]
    except KeyError as exc:
        raise KeyError(f"Latitud no soportada: {latitude_band}") from exc
    try:
        return alt_map[altitude_band]
    except KeyError as exc:
        raise KeyError(f"Altitud no soportada: {altitude_band}") from exc


def _linear_drop(value: float, start: float, end: float) -> float:
    if value < start:
        return 1.0
    if value > end:
        return 0.0
    return (end - value) / (end - start)


def calculate_roof_snow_load(
    latitude_band: str,
    altitude_band: str,
    thermal_condition: str,
    importance_category: str,
    exposure_category: str,
    exposure_condition: str,
    surface_type: str,
    roof_pitch_deg: float,
) -> Dict[str, float]:
    pg_raw = lookup_snow_load(latitude_band, altitude_band)
    pg = _safe_float(pg_raw)
    if pg is None:
        raise ValueError("No existen datos tabulados para la combinación seleccionada.")

    if thermal_condition not in THERMAL_FACTORS:
        raise KeyError(f"Condición térmica no soportada: {thermal_condition}")
    ct = THERMAL_FACTORS[thermal_condition]

    if importance_category not in IMPORTANCE_FACTORS:
        raise KeyError(f"Categoría de importancia no soportada: {importance_category}")
    I = IMPORTANCE_FACTORS[importance_category]

    exposure_map = EXPOSURE_FACTORS.get(exposure_category, EXPOSURE_FACTORS["Otra"])
    if exposure_condition not in exposure_map:
        raise KeyError(f"Condición de exposición no soportada: {exposure_condition}")
    ce = exposure_map[exposure_condition]

    if surface_type not in SURFACE_OPTIONS:
        raise KeyError(f"Tipo de superficie no soportado: {surface_type}")

    slope = roof_pitch_deg
    if surface_type == "Superficie lisas sin obstáculos":
        if ct <= 1.0:
            cs = _linear_drop(slope, 5, 70)
        elif math.isclose(ct, 1.1):
            cs = _linear_drop(slope, 10, 70)
        else:
            cs = _linear_drop(slope, 15, 70)
    else:
        if ct <= 1.0:
            cs = _linear_drop(slope, 30, 70)
        elif math.isclose(ct, 1.1):
            cs = _linear_drop(slope, 39.5, 70)
        else:
            cs = _linear_drop(slope, 45, 70)

    pf = 0.7 * ce * ct * cs * I * pg
    return {"pg": pg, "ct": ct, "I": I, "ce": ce, "cs": round(cs, 3), "pf": round(pf, 3)}


def calculate_seismic_base(
    category: str,
    zone: str,
    soil: str,
    rs_value: float,
    ps_value: float,
    tx: float,
    ty: float,
    r0: float,
    story_heights: Iterable[float],
    story_weights: Iterable[float],
) -> Dict[str, object]:
    if rs_value <= 0 or ps_value <= 0 or tx <= 0 or ty <= 0 or r0 <= 0:
        raise ValueError("Los parámetros sísmicos deben ser mayores que cero.")

    if category == "Categoría I":
        I_s = 0.6
    elif category == "Categoría II":
        I_s = 1.0
    else:
        I_s = 1.2

    if zone == "1":
        A_0 = 0.2
    elif zone == "2":
        A_0 = 0.3
    else:
        A_0 = 0.4

    soil_params = SEISMIC_SOIL.get(soil)
    if not soil_params:
        raise KeyError(f"Tipo de suelo no soportado: {soil}")

    S_S = soil_params["S_S"]
    T0_S = soil_params["T0_S"]
    TP_S = soil_params["TP_S"]
    n_S = soil_params["n_S"]
    p_S = soil_params["p_S"]

    if math.isclose(rs_value, 2.0):
        C_max = 0.9 * S_S * A_0
    elif math.isclose(rs_value, 3.0):
        C_max = 0.6 * S_S * A_0
    elif math.isclose(rs_value, 4.0):
        C_max = 0.55 * S_S * A_0
    elif math.isclose(rs_value, 5.5):
        C_max = 0.4 * S_S * A_0
    else:
        C_max = 0.35 * S_S * A_0

    Co_Sx = (2.75 * S_S * A_0 * (TP_S / tx) ** n_S) / rs_value
    Co_Sy = (2.75 * S_S * A_0 * (TP_S / ty) ** n_S) / rs_value
    C_min = A_0 * S_S / 6
    Q0x = Co_Sx * I_s * ps_value
    Q0y = Co_Sy * I_s * ps_value
    Q0_min = C_min * I_s * ps_value
    Q0_max = C_max * I_s * ps_value
    Qbasx = max(min(Q0x, Q0_max), Q0_min)
    Qbasy = max(min(Q0y, Q0_max), Q0_min)

    Rastx = 1 + tx / (0.1 * T0_S + tx / r0)
    Rasty = 1 + ty / (0.1 * T0_S + ty / r0)

    spectrum = []
    # Generar 21 puntos del espectro: de 0 a 5 segundos con paso de 0.25s
    for i in range(21):
        tn = i * 0.25
        alfa_S = (1 + 4.5 * (tn / T0_S) ** p_S) / (1 + (tn / T0_S) ** 3) if tn > 0 else 1.0
        Sax = S_S * A_0 * alfa_S / (Rastx / I_s)
        SAy = S_S * A_0 * alfa_S / (Rasty / I_s)
        spectrum.append({"period": round(tn, 2), "SaX": round(Sax, 4), "SaY": round(SAy, 4)})

    heights = [float(h) for h in story_heights]
    weights = [float(w) for w in story_weights]
    if len(heights) != len(weights):
        raise ValueError("Las listas de alturas y pesos deben tener la misma longitud.")
    if any(h <= 0 for h in heights) or any(w <= 0 for w in weights):
        raise ValueError("Las alturas y los pesos por nivel deben ser mayores que cero.")

    cumulative_heights = [0.0]
    for value in heights:
        cumulative_heights.append(cumulative_heights[-1] + value)

    H_total = cumulative_heights[-1]
    if H_total <= 0:
        raise ValueError("La suma de alturas debe ser mayor que cero.")

    Ak_pks = 0.0
    for idx in range(1, len(cumulative_heights)):
        upper = cumulative_heights[idx]
        lower = cumulative_heights[idx - 1]

        # Calcular los términos dentro de la raíz
        term_lower = 1.0 - (lower / H_total)
        term_upper = 1.0 - (upper / H_total)

        # Validar que los términos sean no negativos
        if term_lower < 0 or term_upper < 0:
            raise ValueError(
                f"Error en cálculo sísmico: valores fuera de rango en nivel {idx}. "
                f"lower={lower:.3f}, upper={upper:.3f}, H_total={H_total:.3f}, "
                f"term_lower={term_lower:.6f}, term_upper={term_upper:.6f}"
            )

        Ak_pks += (math.sqrt(term_lower) - math.sqrt(term_upper)) * weights[idx - 1]

    if Ak_pks <= 0:
        raise ValueError("La combinación de alturas y pesos genera una base nula; revisa los datos ingresados.")

    floor_forces = []
    for idx in range(1, len(cumulative_heights)):
        upper = cumulative_heights[idx]
        lower = cumulative_heights[idx - 1]

        # Calcular los términos dentro de la raíz
        term_lower = 1.0 - (lower / H_total)
        term_upper = 1.0 - (upper / H_total)

        # Validar que los términos sean no negativos
        if term_lower < 0 or term_upper < 0:
            raise ValueError(
                f"Error en cálculo de fuerzas: valores fuera de rango en nivel {idx}. "
                f"lower={lower:.3f}, upper={upper:.3f}, H_total={H_total:.3f}, "
                f"term_lower={term_lower:.6f}, term_upper={term_upper:.6f}"
            )

        Ak = math.sqrt(term_lower) - math.sqrt(term_upper)
        Fkx = Ak * weights[idx - 1] * Qbasx / Ak_pks
        Fky = Ak * weights[idx - 1] * Qbasy / Ak_pks
        floor_forces.append({"level": idx, "Fkx": round(Fkx, 3), "Fky": round(Fky, 3)})

    return {
        "intensity_factor": I_s,
        "zone_factor": A_0,
        "soil": soil_params,
        "C_max": round(C_max, 4),
        "C_min": round(C_min, 4),
        "Q0x": round(Q0x, 4),
        "Q0y": round(Q0y, 4),
        "Q0_min": round(Q0_min, 4),
        "Q0_max": round(Q0_max, 4),
        "Qbasx": round(Qbasx, 4),
        "Qbasy": round(Qbasy, 4),
        "spectrum": spectrum,
        "floor_forces": floor_forces,
    }


def get_design_base_options() -> Dict[str, object]:
    return {
        "liveLoadCategories": {
            category: [entry["usage"] for entry in entries]
            for category, entries in LIVE_LOADS.items()
        },
        "liveLoadElementTypes": list(KLL_FACTORS.keys()),
        "windEnvironments": list(WIND_PROFILES.keys()),
        "snowLatitudeBands": SNOW_LOAD_MAP,
        "snowThermalConditions": list(THERMAL_FACTORS.keys()),
        "snowImportanceCategories": list(IMPORTANCE_FACTORS.keys()),
        "snowExposureCategories": {
            category: list(options.keys()) for category, options in EXPOSURE_FACTORS.items()
        },
        "snowSurfaceTypes": list(SURFACE_OPTIONS),
        "seismicCategories": ["Categoría I", "Categoría II", "Categoría III"],
        "seismicZones": ["1", "2", "3"],
        "seismicSoils": list(SEISMIC_SOIL.keys()),
    }


def export_design_bases(payload: DesignBaseExportPayload, file_format: str) -> Tuple[bytes, str, str]:
    if file_format == "csv":
        return _export_csv(payload)
    if file_format == "docx":
        return _export_docx(payload)
    if file_format == "pdf":
        return _export_pdf(payload)
    raise ValueError("Formato de exportación no soportado.")


def _export_csv(payload: DesignBaseExportPayload) -> Tuple[bytes, str, str]:
    buffer = io.StringIO()
    writer = csv.writer(buffer)

    if payload.live_load:
        writer.writerow(["Cargas vivas"])
        writer.writerow(["Tipo de edificio", payload.live_load.building_type])
        writer.writerow(["Uso", payload.live_load.usage])
        writer.writerow(
            [
                "Carga uniforme (kN/m²)",
                payload.live_load.uniform_load
                if payload.live_load.uniform_load is not None
                else payload.live_load.uniform_load_raw,
            ]
        )
        writer.writerow(
            [
                "Carga concentrada (kN)",
                payload.live_load.concentrated_load
                if payload.live_load.concentrated_load is not None
                else payload.live_load.concentrated_load_raw,
            ]
        )
        writer.writerow([])

    if payload.reduction:
        writer.writerow(["Reducción por área tributaria"])
        writer.writerow(["Elemento", payload.reduction.element_type])
        writer.writerow(["Área tributaria (m²)", payload.reduction.tributary_area])
        writer.writerow(["Carga base (kN/m²)", payload.reduction.base_load])
        writer.writerow(["Carga reducida (kN/m²)", payload.reduction.reduced_load])
        writer.writerow([])

    if payload.wind:
        writer.writerow(["Presión de viento"])
        writer.writerow(["Entorno", payload.wind.environment])
        writer.writerow(["Altura (m)", payload.wind.height])
        writer.writerow(["q (kN/m²)", payload.wind.q if payload.wind.q is not None else payload.wind.message or "S/I"])
        writer.writerow([])

    if payload.snow:
        writer.writerow(["Cargas de nieve sobre techo"])
        writer.writerow(["Latitud", payload.snow.latitude_band])
        writer.writerow(["Altitud", payload.snow.altitude_band])
        writer.writerow(["Condición térmica", payload.snow.thermal_condition])
        writer.writerow(["Categoría de importancia", payload.snow.importance_category])
        writer.writerow(["Categoría de exposición", payload.snow.exposure_category])
        writer.writerow(["Condición de exposición", payload.snow.exposure_condition])
        writer.writerow(["Tipo de superficie", payload.snow.surface_type])
        writer.writerow(["Inclinación (°)", payload.snow.roof_pitch])
        writer.writerow(["Pg (kN/m²)", payload.snow.pg])
        writer.writerow(["ct", payload.snow.ct])
        writer.writerow(["ce", payload.snow.ce])
        writer.writerow(["I", payload.snow.I])
        writer.writerow(["cs", payload.snow.cs])
        writer.writerow(["pf (kN/m²)", payload.snow.pf])
        writer.writerow([])

    if payload.seismic:
        writer.writerow(["Análisis sísmico base"])
        writer.writerow(["Categoría estructural", payload.seismic.params.category])
        writer.writerow(["Zona", payload.seismic.params.zone])
        writer.writerow(["Suelo", payload.seismic.params.soil])
        writer.writerow(["R", payload.seismic.params.rs])
        writer.writerow(["Peso sísmico (kN)", payload.seismic.params.ps])
        writer.writerow(["Tx (s)", payload.seismic.params.tx])
        writer.writerow(["Ty (s)", payload.seismic.params.ty])
        writer.writerow(["R0", payload.seismic.params.r0])
        writer.writerow(["I_s", payload.seismic.result.intensity_factor])
        writer.writerow(["A0", payload.seismic.result.zone_factor])
        writer.writerow(["Qbas,x (kN)", payload.seismic.result.Qbasx])
        writer.writerow(["Qbas,y (kN)", payload.seismic.result.Qbasy])
        writer.writerow([])
        writer.writerow(["Distribución de fuerzas"])
        writer.writerow(["Nivel", "Fkx (kN)", "Fky (kN)"])
        for row in payload.seismic.result.floor_forces:
            writer.writerow([row.level, row.Fkx, row.Fky])
        writer.writerow([])
        writer.writerow(["Espectro de diseño"])
        writer.writerow(["Periodo (s)", "SaX (g)", "SaY (g)"])
        for point in payload.seismic.result.spectrum:
            writer.writerow([point.period, point.SaX, point.SaY])

    data = buffer.getvalue().encode("utf-8")
    return data, "text/csv; charset=utf-8", "bases_calculo.csv"


def _export_docx(payload: DesignBaseExportPayload) -> Tuple[bytes, str, str]:
    document = Document()
    document.add_heading("Resumen de bases de cálculo", 0)

    if payload.live_load:
        document.add_heading("Cargas vivas", level=1)
        table = document.add_table(rows=4, cols=2)
        table.style = "Light List"
        rows = [
            ("Tipo de edificio", payload.live_load.building_type),
            ("Uso", payload.live_load.usage),
            (
                "Carga uniforme (kN/m²)",
                payload.live_load.uniform_load
                if payload.live_load.uniform_load is not None
                else payload.live_load.uniform_load_raw,
            ),
            (
                "Carga concentrada (kN)",
                payload.live_load.concentrated_load
                if payload.live_load.concentrated_load is not None
                else payload.live_load.concentrated_load_raw,
            ),
        ]
        for idx, (label, value) in enumerate(rows):
            table.cell(idx, 0).text = label
            table.cell(idx, 1).text = str(value)

    if payload.reduction:
        document.add_heading("Reducción por área tributaria", level=1)
        table = document.add_table(rows=4, cols=2)
        table.style = "Light List"
        rows = [
            ("Elemento", payload.reduction.element_type),
            ("Área tributaria (m²)", payload.reduction.tributary_area),
            ("Carga base (kN/m²)", payload.reduction.base_load),
            ("Carga reducida (kN/m²)", payload.reduction.reduced_load),
        ]
        for idx, (label, value) in enumerate(rows):
            table.cell(idx, 0).text = label
            table.cell(idx, 1).text = str(value)

    if payload.wind:
        document.add_heading("Presión de viento", level=1)
        table = document.add_table(rows=3, cols=2)
        table.style = "Light List"
        rows = [
            ("Entorno", payload.wind.environment),
            ("Altura (m)", payload.wind.height),
            (
                "q (kN/m²)",
                payload.wind.q if payload.wind.q is not None else payload.wind.message or "S/I",
            ),
        ]
        for idx, (label, value) in enumerate(rows):
            table.cell(idx, 0).text = label
            table.cell(idx, 1).text = str(value)

    if payload.snow:
        document.add_heading("Cargas de nieve sobre techo", level=1)
        table = document.add_table(rows=11, cols=2)
        table.style = "Light List"
        rows = [
            ("Latitud", payload.snow.latitude_band),
            ("Altitud", payload.snow.altitude_band),
            ("Condición térmica", payload.snow.thermal_condition),
            ("Categoría de importancia", payload.snow.importance_category),
            ("Categoría de exposición", payload.snow.exposure_category),
            ("Condición de exposición", payload.snow.exposure_condition),
            ("Tipo de superficie", payload.snow.surface_type),
            ("Inclinación (°)", payload.snow.roof_pitch),
            ("Pg (kN/m²)", payload.snow.pg),
            ("ct", payload.snow.ct),
            ("pf (kN/m²)", payload.snow.pf),
        ]
        for idx, (label, value) in enumerate(rows):
            table.cell(idx, 0).text = label
            table.cell(idx, 1).text = str(value)

    if payload.seismic:
        document.add_heading("Análisis sísmico base", level=1)
        table = document.add_table(rows=8, cols=2)
        table.style = "Light List"
        rows = [
            ("Categoría estructural", payload.seismic.params.category),
            ("Zona", payload.seismic.params.zone),
            ("Tipo de suelo", payload.seismic.params.soil),
            ("R", payload.seismic.params.rs),
            ("Peso sísmico (kN)", payload.seismic.params.ps),
            ("Tx (s)", payload.seismic.params.tx),
            ("Ty (s)", payload.seismic.params.ty),
            ("R0", payload.seismic.params.r0),
        ]
        for idx, (label, value) in enumerate(rows):
            table.cell(idx, 0).text = label
            table.cell(idx, 1).text = str(value)

        document.add_heading("Distribución de fuerzas por nivel", level=2)
        force_table = document.add_table(rows=len(payload.seismic.result.floor_forces) + 1, cols=3)
        force_table.style = "Light List"
        headers = ["Nivel", "Fkx (kN)", "Fky (kN)"]
        for col, text in enumerate(headers):
            force_table.cell(0, col).text = text
        for idx, row in enumerate(payload.seismic.result.floor_forces, start=1):
            force_table.cell(idx, 0).text = str(row.level)
            force_table.cell(idx, 1).text = f"{row.Fkx:.3f}"
            force_table.cell(idx, 2).text = f"{row.Fky:.3f}"

        document.add_heading("Espectro de diseño (Sa)", level=2)
        spec_table = document.add_table(rows=len(payload.seismic.result.spectrum) + 1, cols=3)
        spec_table.style = "Light List"
        headers = ["Periodo (s)", "SaX (g)", "SaY (g)"]
        for col, text in enumerate(headers):
            spec_table.cell(0, col).text = text
        for idx, point in enumerate(payload.seismic.result.spectrum, start=1):
            spec_table.cell(idx, 0).text = f"{point.period:.2f}"
            spec_table.cell(idx, 1).text = f"{point.SaX:.3f}"
            spec_table.cell(idx, 2).text = f"{point.SaY:.3f}"

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "bases_calculo.docx"


def _export_pdf(payload: DesignBaseExportPayload) -> Tuple[bytes, str, str]:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    x_margin = 40
    y = height - 50

    def write_line(text: str, step: int = 14):
        nonlocal y
        if y < 60:
            pdf.showPage()
            y = height - 50
        pdf.drawString(x_margin, y, text)
        y -= step

    pdf.setTitle("Bases de cálculo")
    pdf.setFont("Helvetica-Bold", 16)
    write_line("Resumen de bases de cálculo", step=22)
    pdf.setFont("Helvetica", 11)

    if payload.live_load:
        pdf.setFont("Helvetica-Bold", 12)
        write_line("Cargas vivas", step=18)
        pdf.setFont("Helvetica", 11)
        write_line(f"Tipo de edificio: {payload.live_load.building_type}")
        write_line(f"Uso: {payload.live_load.usage}")
        uniform_value = (
            payload.live_load.uniform_load
            if payload.live_load.uniform_load is not None
            else payload.live_load.uniform_load_raw
        )
        conc_value = (
            payload.live_load.concentrated_load
            if payload.live_load.concentrated_load is not None
            else payload.live_load.concentrated_load_raw
        )
        write_line(f"Carga uniforme (kN/m²): {uniform_value}")
        write_line(f"Carga concentrada (kN): {conc_value}")
        write_line("")

    if payload.reduction:
        pdf.setFont("Helvetica-Bold", 12)
        write_line("Reducción por área tributaria", step=18)
        pdf.setFont("Helvetica", 11)
        write_line(f"Elemento: {payload.reduction.element_type}")
        write_line(f"Área tributaria (m²): {payload.reduction.tributary_area}")
        write_line(f"Carga base (kN/m²): {payload.reduction.base_load}")
        write_line(f"Carga reducida (kN/m²): {payload.reduction.reduced_load}")
        write_line("")

    if payload.wind:
        pdf.setFont("Helvetica-Bold", 12)
        write_line("Presión de viento", step=18)
        pdf.setFont("Helvetica", 11)
        write_line(f"Entorno: {payload.wind.environment}")
        write_line(f"Altura (m): {payload.wind.height}")
        q_value = payload.wind.q if payload.wind.q is not None else payload.wind.message or "S/I"
        write_line(f"q (kN/m²): {q_value}")
        write_line("")

    if payload.snow:
        pdf.setFont("Helvetica-Bold", 12)
        write_line("Cargas de nieve sobre techo", step=18)
        pdf.setFont("Helvetica", 11)
        entries = [
            ("Latitud", payload.snow.latitude_band),
            ("Altitud", payload.snow.altitude_band),
            ("Condición térmica", payload.snow.thermal_condition),
            ("Categoría de importancia", payload.snow.importance_category),
            ("Categoría de exposición", payload.snow.exposure_category),
            ("Condición de exposición", payload.snow.exposure_condition),
            ("Tipo de superficie", payload.snow.surface_type),
            ("Inclinación (°)", payload.snow.roof_pitch),
            ("Pg (kN/m²)", payload.snow.pg),
            ("ct", payload.snow.ct),
            ("ce", payload.snow.ce),
            ("I", payload.snow.I),
            ("cs", payload.snow.cs),
            ("pf (kN/m²)", payload.snow.pf),
        ]
        for label, value in entries:
            write_line(f"{label}: {value}")
        write_line("")

    if payload.seismic:
        pdf.setFont("Helvetica-Bold", 12)
        write_line("Análisis sísmico base", step=18)
        pdf.setFont("Helvetica", 11)
        entries = [
            ("Categoría estructural", payload.seismic.params.category),
            ("Zona", payload.seismic.params.zone),
            ("Tipo de suelo", payload.seismic.params.soil),
            ("R", payload.seismic.params.rs),
            ("Peso sísmico (kN)", payload.seismic.params.ps),
            ("Tx (s)", payload.seismic.params.tx),
            ("Ty (s)", payload.seismic.params.ty),
            ("R0", payload.seismic.params.r0),
            ("I_s", payload.seismic.result.intensity_factor),
            ("A0", payload.seismic.result.zone_factor),
            ("Qbas,x (kN)", payload.seismic.result.Qbasx),
            ("Qbas,y (kN)", payload.seismic.result.Qbasy),
        ]
        for label, value in entries:
            write_line(f"{label}: {value}")
        write_line("")

        pdf.setFont("Helvetica-Bold", 11)
        write_line("Distribución de fuerzas (kN)", step=16)
        pdf.setFont("Helvetica", 10)
        write_line("Nivel    Fkx      Fky")
        for row in payload.seismic.result.floor_forces:
            write_line(f"{row.level:>3}   {row.Fkx:>7.2f}   {row.Fky:>7.2f}")
        write_line("")

        pdf.setFont("Helvetica-Bold", 11)
        write_line("Espectro de diseño Sa (g)", step=16)
        pdf.setFont("Helvetica", 10)
        write_line("Periodo    SaX    SaY")
        for point in payload.seismic.result.spectrum:
            write_line(f"{point.period:>6.2f}   {point.SaX:>5.3f}   {point.SaY:>5.3f}")

    pdf.showPage()
    pdf.save()
    return buffer.getvalue(), "application/pdf", "bases_calculo.pdf"
